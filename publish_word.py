import re
from docx import Document
from docx.oxml.ns import qn
from urllib.parse import quote
from backend_api import IShareUploader, DEFAULT_CONFIG

class WordUploader(IShareUploader):
    def __init__(self, monthly_post_id, config=None):
        # 合併 Config
        final_config = config or DEFAULT_CONFIG.copy()
        final_config['MONTHLY_POST_ID'] = str(monthly_post_id)
        super().__init__(config=final_config)
    
    def extract_table_html(self, table):
        """
        將 Word 表格轉換為 Mso-style HTML
        
        Args:
            table: python-docx 表格對象
            
        Returns:
            str: Mso-style HTML 格式的表格
        """
        html_parts = []
        
        # 開始表格標籤
        html_parts.append(
            '<table class="MsoNormalTable" border="1" cellspacing="0" cellpadding="0" '
            'style="border-collapse:collapse;border:none;mso-border-alt:solid windowtext .5pt;'
            'mso-yfti-tbllook:1184;mso-padding-alt:0cm 5.4pt 0cm 5.4pt">'
        )
        
        # 遍歷表格行
        for row_idx, row in enumerate(table.rows):
            # 行樣式（標記首行）
            row_style = 'mso-yfti-irow:0;mso-yfti-firstrow:yes' if row_idx == 0 else f'mso-yfti-irow:{row_idx}'
            html_parts.append(f'  <tr style="{row_style}">')
            
            # 遍歷儲存格
            for cell in row.cells:
                # 提取儲存格背景色
                cell_xml = cell._element
                bg_color = None
                if cell_xml.tcPr is not None:
                    shading = cell_xml.tcPr.find(qn('w:shd'))
                    if shading is not None:
                        fill = shading.get(qn('w:fill'))
                        if fill and fill != 'auto':
                            bg_color = f'#{fill}'
                
                # 建立儲存格樣式
                cell_style_parts = [
                    'border:solid windowtext 1.0pt',
                    'mso-border-alt:solid windowtext .5pt',
                    'padding:0cm 5.4pt 0cm 5.4pt'
                ]
                if bg_color:
                    cell_style_parts.append(f'background:{bg_color}')
                
                cell_style = ';'.join(cell_style_parts)
                html_parts.append(f'    <td valign="top" style="{cell_style}">')
                
                # 處理儲存格內的段落
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if not para_text:
                        html_parts.append('      <p class="MsoNormal"><br></p>')
                        continue
                    
                    # 檢查段落樣式
                    is_bold = any(run.bold for run in para.runs if run.bold is not None)
                    
                    # 提取顏色和字級
                    color = None
                    font_size = None
                    for run in para.runs:
                        if run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            color = f"#{rgb}"
                        if run.font.size:
                            font_size = run.font.size.pt
                        if color and font_size:
                            break
                    
                    # 建立 span 樣式
                    style_parts = ['font-family:&quot;微軟正黑體&quot;,sans-serif']
                    if color:
                        style_parts.append(f'color:{color}')
                    if font_size:
                        style_parts.append(f'font-size:{font_size}pt')
                    span_style = ';'.join(style_parts)
                    
                    # 生成 HTML
                    if is_bold:
                        html_parts.append(
                            f'      <p class="MsoNormal"><b><span style="{span_style}">{para_text}</span></b></p>'
                        )
                    else:
                        html_parts.append(
                            f'      <p class="MsoNormal"><span style="{span_style}">{para_text}</span></p>'
                        )
                
                html_parts.append('    </td>')
            
            html_parts.append('  </tr>')
        
        # 結束表格標籤
        html_parts.append('</table>')
        
        return '\n'.join(html_parts)

    def process_docx(self, file_stream):
        """
        處理 Word 文件串流 (BytesIO)，回傳解析後的 sections
        支援段落、圖片、表格的混合處理
        """
        doc = Document(file_stream)
        sections = []
        current_html_parts = []
        
        # 建立 relationship ID 到圖片數據的映射
        image_rels = {}
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_rels[rel.rId] = rel.target_part.blob
        
        # 建立文件元素索引（包含段落和表格）以保持正確順序
        body_elements = []
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                body_elements.append(('paragraph', element))
            elif element.tag.endswith('tbl'):  # 表格
                body_elements.append(('table', element))
        
        # 依序處理所有元素
        para_index = 0
        table_index = 0
        
        for elem_type, elem in body_elements:
            if elem_type == 'paragraph':
                para = doc.paragraphs[para_index]
                para_index += 1
                
                # 檢查段落中是否有圖片
                has_image = False
                for run in para.runs:
                    # 檢查 run 中的 drawing 元素
                    drawings = run._element.findall('.//' + qn('a:blip'))
                    for blip in drawings:
                        embed_id = blip.get(qn('r:embed'))
                        if embed_id and embed_id in image_rels:
                            has_image = True
                            # 先保存之前累積的文字 HTML
                            if current_html_parts:
                                html_text = '\n'.join(current_html_parts)
                                if html_text:
                                    sections.append({'type': 'text', 'content': html_text})
                                current_html_parts = []
                            # 保存圖片 (bytes)
                            sections.append({'type': 'image', 'content': image_rels[embed_id]})
                
                # 如果段落沒有圖片，處理文字格式
                if not has_image:
                    text = para.text.strip()
                    if not text:
                        # 空行
                        current_html_parts.append('<p class="MsoNormal"><br></p>')
                        continue

                    # 取得段落樣式資訊
                    style_name = para.style.name if para.style else ""
                    is_heading = 'Heading' in style_name or 'Title' in style_name
                    
                    # 更精確的列點辨識邏輯
                    # 1. 檢查段落樣式名稱
                    is_list = 'List' in style_name
                    
                    # 2. 檢查段落的 numbering 屬性（Word 內部列表標記）
                    if not is_list and para._element.pPr is not None:
                        numPr = para._element.pPr.find(qn('w:numPr'))
                        if numPr is not None:
                            is_list = True
                    
                    # 3. 檢查文字是否以常見列點符號開頭
                    if not is_list:
                        is_list = text.startswith('•') or text.startswith('-') or text.startswith('‧') or re.match(r'^\d+\.', text) or text.startswith('◆') or text.startswith('▪')
                    
                    # 檢查粗體
                    has_bold = any(run.bold for run in para.runs if run.bold is not None)
                    
                    # 檢查顏色與字級資訊
                    color = None
                    font_size = None
                    for run in para.runs:
                        # 抓取顏色
                        if run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            color = f"#{rgb}"
                        # 抓取字級 (Word 使用 Twip 單位，需轉換為 pt)
                        if run.font.size:
                            # Twip to Point: 1 pt = 20 twips
                            font_size = run.font.size.pt
                        # 如果已取得顏色和字級則跳出
                        if color and font_size:
                            break
                    
                    # 建立 inline style
                    style_parts = ['font-family:&quot;微軟正黑體&quot;,sans-serif;']
                    if color:
                        style_parts.append(f'color:{color};')
                    if font_size:
                        style_parts.append(f'font-size:{font_size}pt;')
                    inline_style = ''.join(style_parts)
                    
                    # 根據樣式產生對應的 HTML
                    if is_heading:
                        # 標題樣式：若沒有明確顏色則使用預設藍色
                        if not color:
                            style_parts_heading = ['font-family:&quot;微軟正黑體&quot;,sans-serif;', 'color:#0070C0;']
                            if font_size:
                                style_parts_heading.append(f'font-size:{font_size}pt;')
                            inline_style = ''.join(style_parts_heading)
                        html = f'<p class="MsoNormal"><b><span style="{inline_style}">{text}</span></b></p>'
                        current_html_parts.append(html)
                    elif has_bold:
                        # 一般粗體文字：保持原有顏色（若 Word 中為黑色則不加 color 屬性）
                        html = f'<p class="MsoNormal"><b><span style="{inline_style}">{text}</span></b></p>'
                        current_html_parts.append(html)
                    elif is_list:
                        # 列表項目 - 移除開頭的列點符號（如果有的話）
                        clean_text = re.sub(r'^[•\-‧◆▪\d+\.]+\s*', '', text)
                        # 使用 Mso-style 列表格式，加入樣式
                        list_style = inline_style if (color or font_size) else 'font-family:&quot;微軟正黑體&quot;,sans-serif;'
                        html = f'<p class="MsoListParagraphCxSpFirst" style="margin-left:24.0pt;mso-add-space:auto;text-indent:-24.0pt;mso-list:l0 level1 lfo1;layout-grid-mode:char"><span style="{list_style}"><span style="font-family:Wingdings;mso-fareast-font-family:新細明體;mso-bidi-font-family:新細明體">l<span style="font:7.0pt &quot;Times New Roman&quot;">&nbsp; </span></span>{clean_text}</span></p>'
                        current_html_parts.append(html)
                    else:
                        # 一般段落（若有字級或顏色才加 span）
                        if font_size or color:
                            html = f'<p class="MsoNormal"><span style="{inline_style}">{text}</span></p>'
                        else:
                            html = f'<p class="MsoNormal">{text}</p>'
                        current_html_parts.append(html)
            
            elif elem_type == 'table':
                # 處理表格
                table = doc.tables[table_index]
                table_index += 1
                
                # 保存之前累積的文字 HTML
                if current_html_parts:
                    html_text = '\n'.join(current_html_parts)
                    if html_text:
                        sections.append({'type': 'text', 'content': html_text})
                    current_html_parts = []
                
                # 生成表格 HTML 並作為 text section
                table_html = self.extract_table_html(table)
                sections.append({'type': 'text', 'content': table_html})
        
        # 保存最後累積的文字 HTML
        if current_html_parts:
            html_text = '\n'.join(current_html_parts)
            if html_text:
                sections.append({'type': 'text', 'content': html_text})
        
        return sections

    def upload_sections(self, sections):
        """上傳解析後的 sections 到 iShare"""
        final_post_data = []
        
        for section in sections:
            if section['type'] == 'image':
                # 上傳圖片 (假設是 jpg)
                img_url = self.upload_image_bytes(section['content'], "uploaded_image.jpg")
                if img_url:
                    payload = self.build_section_payload(1, "", "", photo_url=img_url, alt="Word Image")
                    final_post_data.append(payload)
            else:
                # 文字
                payload = self.build_section_payload(8, "", section['content'])
                final_post_data.append(payload)
        
        if final_post_data:
            self.submit_data(final_post_data)
            return True, f"成功上傳 {len(final_post_data)} 個段落"
        else:
            return False, "沒有可上傳的資料"

# 相容舊版介面 (如果需要保留直接執行能力，可以保留下方區塊，否則可移除)
if __name__ == "__main__":
    print("This module is now a library. Please use WordUploader class.")
